"""
MDRF Quarterly Report Template v4 — Free-form style (matching MLRF kit)
Each section: banner → metadata → open writing area with guidance → sources → signatures
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

NAVY = '1B2A4A'; BLUE = '2563EB'; GREEN = '059669'
AMBER = 'D97706'; RED = 'DC2626'; LT = 'F8FAFC'; GREY = '6B7280'

DISTRICTS = [
    'East Khasi Hills', 'West Khasi Hills', 'South West Khasi Hills',
    'Ri-Bhoi', 'Jaintia Hills', 'East Jaintia Hills', 'West Jaintia Hills',
    'East Garo Hills', 'West Garo Hills', 'South Garo Hills',
    'North Garo Hills', 'South West Garo Hills',
]

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def col(table, idx, cm):
    for row in table.rows:
        row.cells[idx].width = Cm(cm)

def tbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    return t

def run(p, text, bold=False, sz=10, color=None, italic=False):
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(sz)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:6],16))
    return r

def banner(doc, code, title, subtitle):
    t = tbl(doc, 1, 2)
    col(t, 0, 11); col(t, 1, 5)
    left = t.cell(0, 0); shade(left, NAVY)
    p = left.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    run(p, 'MEGHALAYA INSTITUTE OF GOVERNANCE', bold=True, sz=9, color='FFFFFF')
    p2 = left.add_paragraph()
    run(p2, 'Meghalaya District Research Fellowship', sz=8, color='93C5FD')
    p3 = left.add_paragraph()
    p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(6)
    run(p3, title, bold=True, sz=16, color='FFFFFF')
    if subtitle:
        p4 = left.add_paragraph()
        run(p4, subtitle, sz=8, color='93C5FD')
    right = t.cell(0, 1); shade(right, BLUE)
    p = right.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    run(p, code, bold=True, sz=22, color='FFFFFF')
    p2 = right.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, 'DOCUMENT ID', bold=True, sz=8, color='DBEAFE')
    p3 = right.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p3, 'MDRF-2026-Q_-XX-###', sz=8, color='DBEAFE')
    p4 = right.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p4, 'Assigned by MIG on receipt', sz=7, color='93C5FD')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def meta(doc, fields):
    t = tbl(doc, len(fields) + 1, 2)
    col(t, 0, 5); col(t, 1, 11)
    for i, (label, placeholder) in enumerate(fields):
        shade(t.cell(i, 0), LT)
        run(t.cell(i, 0).paragraphs[0], label, bold=True, sz=10)
        run(t.cell(i, 1).paragraphs[0], placeholder, sz=9, color=GREY)
    last = len(fields)
    shade(t.cell(last, 0), LT)
    run(t.cell(last, 0).paragraphs[0], 'OFFICIAL USE ONLY', bold=True, sz=9, color=RED)
    run(t.cell(last, 1).paragraphs[0],
        'Reviewed by: ___________________   Date: __________   ☐ Approved  ☐ Revision', sz=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def freeform(doc, guidance, lines=20):
    """MLRF-style: single cell with guidance text + empty lines for writing."""
    t = tbl(doc, 1, 1)
    c = t.cell(0, 0)
    p = c.paragraphs[0]
    # Guidance in grey italic
    run(p, '(', sz=10, color=GREY)
    run(p, guidance, sz=10, color=GREY, italic=True)
    run(p, ')', sz=10, color=GREY)
    # Empty writing lines
    for _ in range(lines):
        c.add_paragraph()
    return t


def freeform_with_hint(doc, hint, lines=15):
    """Free-form with a bold hint at the top."""
    t = tbl(doc, 1, 1)
    c = t.cell(0, 0)
    p = c.paragraphs[0]
    run(p, hint, bold=True, sz=10)
    for _ in range(lines):
        c.add_paragraph()
    return t


def sigs(doc):
    t = tbl(doc, 1, 2)
    for j, label in enumerate(['MDRF Fellow', 'State Coordinator / SRF']):
        c = t.cell(0, j)
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(6)
        run(p, label, bold=True, sz=9, color=NAVY)
        for _ in range(3): c.add_paragraph()
        run(c.add_paragraph(), 'Name: ______________________________', sz=9)
        run(c.add_paragraph(), 'Date: ______________________________', sz=9)


def sources(doc, rows=5):
    t = tbl(doc, rows + 1, 3)
    col(t, 0, 1.5); col(t, 1, 8); col(t, 2, 6.5)
    for j, h in enumerate(['Sl.', 'Source / Document', 'URL / Details']):
        shade(t.cell(0, j), NAVY)
        run(t.cell(0, j).paragraphs[0], h, bold=True, sz=9, color='FFFFFF')
    for i in range(1, rows + 1):
        run(t.cell(i, 0).paragraphs[0], str(i), sz=9)


def page_break(doc):
    doc.add_page_break()


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# MAIN
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def create():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── TITLE PAGE ───────────────────────────────────────────────
    for _ in range(3): doc.add_paragraph()
    t = tbl(doc, 1, 1)
    c = t.cell(0, 0); shade(c, NAVY)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run(p, 'MEGHALAYA INSTITUTE OF GOVERNANCE', bold=True, sz=14, color='FFFFFF')
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, 'Meghalaya District Research Fellowship', sz=11, color='93C5FD')
    p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(16)
    run(p3, 'FELLOW QUARTERLY REPORT', bold=True, sz=26, color='FFFFFF')
    p4 = c.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(20)
    run(p4, 'District Research Fellows', bold=True, sz=16, color='34D399')

    doc.add_paragraph()
    t2 = tbl(doc, 1, 1)
    shade(t2.cell(0, 0), 'EFF6FF')
    p = t2.cell(0, 0).paragraphs[0]
    run(p, 'HOW TO USE — 3 STEPS', bold=True, sz=11)
    p2 = t2.cell(0, 0).add_paragraph()
    run(p2, 'Step 1', bold=True, sz=10, color=GREEN)
    run(p2, '  Fill the COVER SHEET — enter all fields. Grey rows are completed by MIG.', sz=10)
    p3 = t2.cell(0, 0).add_paragraph()
    run(p3, 'Step 2', bold=True, sz=10, color=GREEN)
    run(p3, '  Write in the BODY — section headings are pre-printed. Write freely under each heading.', sz=10)
    p4 = t2.cell(0, 0).add_paragraph()
    run(p4, 'Step 3', bold=True, sz=10, color=GREEN)
    run(p4, '  Submit to MIG — attach evidence (photos, certificates) as appendices.', sz=10)

    doc.add_paragraph()
    t3 = tbl(doc, 8, 3)
    col(t3, 0, 2); col(t3, 1, 5); col(t3, 2, 9)
    for j, h in enumerate(['CODE', 'SECTION', 'USE WHEN']):
        shade(t3.cell(0, j), NAVY)
        run(t3.cell(0, j).paragraphs[0], h, bold=True, sz=9, color='FFFFFF')
    for i, (code, name, desc) in enumerate([
        ('CP',  'Cover Page',             'Always — fellow details, district, quarter'),
        ('WA',  'Work Summary',           'Always — overview of all work done this quarter'),
        ('FV',  'Field Visits',           'If you did fieldwork with the DC'),
        ('DR',  'Data & Research',        'If you collected data or produced research outputs'),
        ('LC',  'LMS & Capacity Building','If you did LMS courses, sessions, residencies, or training'),
        ('SU',  'Challenges & Support',   'If you faced issues or need support'),
        ('NP',  'Next Quarter Plan',      'Always — what you plan to do next'),
    ]):
        run(t3.cell(i+1, 0).paragraphs[0], code, sz=9)
        run(t3.cell(i+1, 1).paragraphs[0], name, sz=9)
        run(t3.cell(i+1, 2).paragraphs[0], desc, sz=9, color=GREY)

    doc.add_paragraph()
    t4 = tbl(doc, 1, 1)
    shade(t4.cell(0, 0), 'EFF6FF')
    p = t4.cell(0, 0).paragraphs[0]
    run(p, 'DOCUMENT ID FORMAT\n', bold=True, sz=10, color=NAVY)
    run(p, 'MDRF-2026-Q_-[CODE]-[DISTRICT]-[INITIALS]-[###]\n', bold=True, sz=10, color=BLUE)
    run(p, 'Example: MDRF-2026-Q1-WA-EKH-PS-001  ·  EKH · WKH · SWKH · RB · JH · EJH · WJH · EGH · WGH · SGH · NGH · SWGH', sz=8, color=GREY)

    page_break(doc)

    # ── CP: COVER PAGE ───────────────────────────────────────────
    banner(doc, 'CP', 'COVER PAGE', 'Fellow details & programme information')

    dist_list = '  ☐ '.join(DISTRICTS)
    meta(doc, [
        ('MDRF Fellow',       '  Your full name'),
        ('District',          f'  ☐ {dist_list}'),
        ('Quarter',           '  ☐ Q1 (Apr–Jun)  ☐ Q2 (Jul–Sep)  ☐ Q3 (Oct–Dec)  ☐ Q4 (Jan–Mar)'),
        ('Reporting Period',  '  e.g. April 1 – June 30, 2026'),
        ('Date Submitted',    ' '),
        ('Programme Year',    '  ☐ Year 1  ☐ Year 2'),
        ('Cohort',            '  ☐ New  ☐ Existing'),
        ('District Collector','  Name of DC'),
        ('Assigned Dept',     '  Planning / Rural / Other'),
        ('State Coordinator', '  Coordinator overseeing you'),
        ('SRF',               '  SRF you work with'),
        ('Data Scientist',    '  Ajay Choudhary'),
    ])

    sigs(doc)
    page_break(doc)

    # ── WA: WORK SUMMARY ─────────────────────────────────────────
    banner(doc, 'WA', 'WORK SUMMARY', 'Overview of all work done this quarter')

    meta(doc, [
        ('MDRF Fellow',  '  Your name'),
        ('District',     ' '),
        ('Quarter',      ' '),
        ('Date Submitted',' '),
        ('Total Working Days', ' '),
        ('Field Visits',  ' '),
        ('Reports / Outputs', ' '),
    ])

    p = doc.add_paragraph()
    run(p, 'Executive Summary', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'Provide a 200–300 word overview of your key activities and achievements this quarter. '
        'Cover: main assignments from the DC, fieldwork done, data collected, reports produced, '
        'LMS courses completed, monthly sessions attended, and any significant outcomes.',
        lines=20)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Month-by-Month Account', bold=True, sz=12, color=NAVY)

    for month in ['Month 1', 'Month 2', 'Month 3']:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run(p, month, bold=True, sz=11, color=BLUE)
        freeform(doc,
            f'Describe your work during {month}. Include: '
            'dates, locations visited, DC assignments completed, data collected, '
            'meetings attended, reports submitted, and any other significant work.',
            lines=8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Key Achievements This Quarter', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'List your top 3–5 achievements or milestones this quarter. '
        'These could be: a report completed, a dataset compiled, a successful field visit, '
        'positive feedback from DC, LMS course completed, blog published, etc.',
        lines=8)

    page_break(doc)

    # ── FV: FIELD VISITS ─────────────────────────────────────────
    banner(doc, 'FV', 'FIELD VISITS', 'District-level fieldwork log')

    meta(doc, [
        ('MDRF Fellow',  '  Your name'),
        ('District',     ' '),
        ('Quarter',      ' '),
        ('Date Submitted',' '),
    ])

    p = doc.add_paragraph()
    run(p, 'Field Visit Log', bold=True, sz=12, color=NAVY)

    # Structured table for field visits
    t = tbl(doc, 11, 6)
    col(t, 0, 1); col(t, 1, 2.5); col(t, 2, 3.5); col(t, 3, 3); col(t, 4, 3); col(t, 5, 3)
    for j, h in enumerate(['Sl.', 'Date', 'Location / Village', 'Block', 'Purpose', 'Key Findings']):
        shade(t.cell(0, j), NAVY)
        run(t.cell(0, j).paragraphs[0], h, bold=True, sz=9, color='FFFFFF')
    for i in range(1, 11):
        run(t.cell(i, 0).paragraphs[0], str(i), sz=8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Field Visit Narrative', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'Write in detail about your most significant field visits. '
        'What did you observe? What data did you collect? '
        'Who did you meet? What were the key findings? '
        'What challenges did you face in the field?',
        lines=15)

    p = doc.add_paragraph()
    run(p, 'Attach field visit photographs as Appendix A.', sz=9, color=GREY)

    page_break(doc)

    # ── DR: DATA & RESEARCH ──────────────────────────────────────
    banner(doc, 'DR', 'DATA & RESEARCH', 'Datasets, analyses, research outputs')

    meta(doc, [
        ('MDRF Fellow',  '  Your name'),
        ('District',     ' '),
        ('Quarter',      ' '),
        ('Date Submitted',' '),
    ])

    p = doc.add_paragraph()
    run(p, 'Data Collection & Compilation', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'Describe the data you collected or compiled this quarter. '
        'What datasets did you work with? What sources did you use? '
        'How many records? What was the data for? '
        'Include any data from India Data Portal (IDP).',
        lines=12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Research Outputs', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'Describe any research outputs: reports, analyses, dashboards, briefs, or data products. '
        'Who were they shared with? What impact did they have? '
        'Include any contributions to CDP or DDP.',
        lines=12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'India Data Portal (IDP) Usage', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'What data did you access from the India Data Portal? '
        'What analyses did you perform? What insights did you generate? '
        'How did IDP support your district work?',
        lines=8)

    sources(doc)
    page_break(doc)

    # ── LC: LMS & CAPACITY BUILDING ──────────────────────────────
    banner(doc, 'LC', 'LMS & CAPACITY BUILDING',
           'ISB LMS · iGOT · Monthly sessions · Residencies · Training')

    meta(doc, [
        ('MDRF Fellow',  '  Your name'),
        ('District',     ' '),
        ('Quarter',      ' '),
        ('Date Submitted',' '),
    ])

    p = doc.add_paragraph()
    run(p, 'LMS Coursework (ISB + iGOT)', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'List the LMS courses you enrolled in or completed this quarter. '
        'Include: course name, platform (ISB LMS / iGOT), status (enrolled / in progress / completed), '
        'and any assessment scores. Courses may cover: climate change, India Data Portal, '
        'team building, leadership, psychometric tools, public policy fundamentals, etc.',
        lines=10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Monthly Live Sessions', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'Record the monthly live online sessions you attended. '
        'Sessions are compulsory, held once a month (e.g., first Wednesday at 5:00 PM). '
        'Include: date, topic/theme, speaker, and key takeaways. '
        'Sessions may cover CDP, DDP, blog writing, capacity building, or feedback from MBDA.',
        lines=10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Residencies', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'If you attended a residency this quarter, describe it. '
        'Residencies are 3–4 days long, held at ISB campus (Hyderabad or Mohali) or Shillong. '
        'Include: location, dates, key sessions/workshops, learnings, and networking. '
        'Four residencies over 2 years: twice at ISB, twice at Shillong.',
        lines=10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Blog Writing & Communication', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'List any blogs, articles, or communication pieces you wrote this quarter. '
        'Blog writing is part of capacity building. Include: title, topic, where published, '
        'and any feedback received.',
        lines=8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Induction / Immersion Programme', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'If you participated in an induction or immersion programme (10 days, Shillong), '
        'describe your experience. Monday–Friday classes, Saturday excursion, Sunday off. '
        'Focus on bonding, knowing fellow fellows, capacity building.',
        lines=8)

    sources(doc)
    page_break(doc)

    # ── SU: CHALLENGES & SUPPORT ─────────────────────────────────
    banner(doc, 'SU', 'CHALLENGES & SUPPORT', 'Issues faced · Escalations · Support needed')

    meta(doc, [
        ('MDRF Fellow',  '  Your name'),
        ('District',     ' '),
        ('Quarter',      ' '),
        ('Date Submitted',' '),
    ])

    p = doc.add_paragraph()
    run(p, '1-on-1 Interaction with Policy Director', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'Record your interaction with the Policy Director. '
        'Share your experiences, success stories, and lessons learned. '
        'Note any supports needed by the DC that you cannot handle on your own. '
        'SRFs and Data Scientist may be engaged to provide support.',
        lines=10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Challenges Faced', bold=True, sz=12, color=RED)

    freeform(doc,
        'Describe the challenges you faced this quarter. '
        'Categories: data access, admin approvals, technical issues, research gaps, '
        'logistics, LMS problems, DC office coordination, or any other. '
        'What was the impact? How did you try to resolve them?',
        lines=12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Support Needed', bold=True, sz=12, color=AMBER)

    freeform(doc,
        'What support do you need from the State Coordinator, SRF, Data Scientist, or MIG? '
        'Be specific: what do you need, from whom, and how urgent is it? '
        'Include any DC support needs that you cannot handle alone.',
        lines=10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'SRF / Data Scientist Support Received', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'What support did you receive from the SRF and Data Scientist this quarter? '
        'Was it helpful? What additional support do you need from them?',
        lines=6)

    page_break(doc)

    # ── NP: NEXT QUARTER PLAN ────────────────────────────────────
    banner(doc, 'NP', 'NEXT QUARTER PLAN', 'Planned activities, learning goals, targets')

    meta(doc, [
        ('MDRF Fellow',  '  Your name'),
        ('District',     ' '),
        ('Quarter',      '  Next quarter'),
        ('Date Submitted',' '),
    ])

    p = doc.add_paragraph()
    run(p, 'Planned Activities', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'What are your key activities planned for next quarter? '
        'Include: DC assignments expected, fieldwork planned, data collection targets, '
        'reports to be produced, CDP/DDP work, LMS courses to complete, '
        'and any upcoming residencies or training.',
        lines=15)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Learning & Development Goals', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'What skills or knowledge do you plan to develop next quarter? '
        'Technical skills (data tools, GIS, etc.), research skills (writing, analysis), '
        'domain knowledge (public policy, climate change, etc.), '
        'and any LMS courses you plan to complete.',
        lines=8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, 'Residency Preparation', bold=True, sz=12, color=NAVY)

    freeform(doc,
        'Is there an upcoming residency? What preparation do you need? '
        'What topics would you like covered? Any logistics support needed?',
        lines=5)

    page_break(doc)

    # ── APPENDIX A ───────────────────────────────────────────────
    p = doc.add_paragraph()
    run(p, 'APPENDIX A: FIELD VISIT PHOTOGRAPHS', bold=True, sz=14, color=NAVY)
    p2 = doc.add_paragraph()
    run(p2, 'Paste field visit photographs below with captions (location, date, purpose).',
        sz=9, color=GREY)

    t = tbl(doc, 4, 2)
    col(t, 0, 8); col(t, 1, 8)
    for i in range(4):
        for j in range(2):
            c = t.cell(i, j); c.height = Cm(4.5)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p, '\n\n[Insert Photo]', sz=10, color='B0B0B0')
            p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p2, 'Location: __________  Date: __________  Purpose: __________', sz=8, color=GREY)

    page_break(doc)

    # ── APPENDIX B ───────────────────────────────────────────────
    p = doc.add_paragraph()
    run(p, 'APPENDIX B: CERTIFICATES & EVIDENCE', bold=True, sz=14, color=NAVY)
    p2 = doc.add_paragraph()
    run(p2, 'Attach LMS completion certificates, residency attendance proof, blog screenshots, etc.',
        sz=9, color=GREY)

    t = tbl(doc, 4, 2)
    col(t, 0, 8); col(t, 1, 8)
    for i in range(4):
        for j in range(2):
            c = t.cell(i, j); c.height = Cm(4.5)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p, '\n[Insert Certificate / Screenshot]', sz=10, color='B0B0B0')
            p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p2, 'Type: __________  Date: __________', sz=8, color=GREY)

    page_break(doc)

    # ── APPENDIX C ───────────────────────────────────────────────
    p = doc.add_paragraph()
    run(p, 'APPENDIX C: SOURCES & REFERENCES', bold=True, sz=14, color=NAVY)
    p2 = doc.add_paragraph()
    run(p2, 'List all sources, references, and data portals used during the quarter.',
        sz=9, color=GREY)

    sources(doc, rows=10)

    doc.add_paragraph()
    doc.add_paragraph()
    sigs(doc)

    # ── SAVE ─────────────────────────────────────────────────────
    fp = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      'MDRF_Quarterly_Report_Template.docx')
    doc.save(fp)
    return fp


if __name__ == '__main__':
    print('Generating MDRF Quarterly Report Template v4...\n')
    fp = create()
    print(f'  ✅ {fp}')
    print(f'\n🎉 Done — free-form style, matching MLRF kit.')
